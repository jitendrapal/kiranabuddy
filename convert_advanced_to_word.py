from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Create a new Document
doc = Document()

# Read the markdown file
print("Reading JAVA_ADVANCED_COMPLETE_BOOK.md...")
with open('JAVA_ADVANCED_COMPLETE_BOOK.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

print(f"Processing {len(lines)} lines...")

in_code_block = False
code_lines = []
current_line = 0

for line in lines:
    current_line += 1
    if current_line % 1000 == 0:
        print(f"Processing line {current_line}/{len(lines)}...")
    
    # Handle code blocks
    if line.strip().startswith('```'):
        if in_code_block:
            # End of code block - add all code lines
            if code_lines:
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Normal'
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 128)
                # Add background color effect
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.right_indent = Inches(0.5)
            code_lines = []
            in_code_block = False
        else:
            in_code_block = True
        continue
    
    if in_code_block:
        code_lines.append(line.rstrip())
        continue
    
    # Skip empty lines
    if not line.strip():
        doc.add_paragraph()
        continue
    
    # Main title
    if line.startswith('# '):
        p = doc.add_heading(line[2:].strip(), level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = RGBColor(0, 51, 102)
        doc.add_page_break()
        continue
    
    # Chapter headings
    if line.startswith('## '):
        doc.add_page_break()
        p = doc.add_heading(line[3:].strip(), level=1)
        for run in p.runs:
            run.font.color.rgb = RGBColor(204, 0, 0)
        continue
    
    # Section headings
    if line.startswith('### '):
        p = doc.add_heading(line[4:].strip(), level=2)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0, 102, 204)
        continue
    
    # Subsection headings
    if line.startswith('#### '):
        p = doc.add_heading(line[5:].strip(), level=3)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0, 153, 76)
        continue
    
    # Bullet points
    if line.strip().startswith('- ') or line.strip().startswith('* '):
        text = line.strip()[2:]
        p = doc.add_paragraph(text, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.5)
        continue
    
    # Numbered lists
    if re.match(r'^\d+\.\s', line.strip()):
        text = re.sub(r'^\d+\.\s', '', line.strip())
        p = doc.add_paragraph(text, style='List Number')
        p.paragraph_format.left_indent = Inches(0.5)
        continue
    
    # Tables (simple markdown tables)
    if '|' in line and not line.strip().startswith('//'):
        # Skip table separator lines
        if re.match(r'^\|[\s\-:]+\|', line.strip()):
            continue
        # This is a table row
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        # For simplicity, just add as formatted text
        p = doc.add_paragraph(' | '.join(cells))
        p.paragraph_format.left_indent = Inches(0.3)
        for run in p.runs:
            run.font.size = Pt(10)
        continue
    
    # Horizontal rules
    if line.strip() == '---':
        p = doc.add_paragraph('_' * 80)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = RGBColor(192, 192, 192)
        continue
    
    # Bold text with **
    if '**' in line:
        p = doc.add_paragraph()
        parts = line.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            if i % 2 == 1:  # Odd indices are bold
                run.bold = True
        continue
    
    # Regular paragraph
    p = doc.add_paragraph(line.strip())
    p.paragraph_format.line_spacing = 1.15

# Save the document
print("\nSaving document as JAVA_ADVANCED_COMPLETE_BOOK.docx...")
doc.save('JAVA_ADVANCED_COMPLETE_BOOK.docx')

# Calculate approximate page count (assuming ~40 lines per page)
approx_pages = len(lines) // 40
print(f"✅ Successfully created JAVA_ADVANCED_COMPLETE_BOOK.docx")
print(f"📄 Total pages: approximately {approx_pages}")

