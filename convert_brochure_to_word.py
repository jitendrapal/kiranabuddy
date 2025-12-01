from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Read the markdown file
print("Reading Neuro_Edge_Technologies_Brochure.md...")
with open('Neuro_Edge_Technologies_Brochure.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

print(f"Processing {len(lines)} lines...")

in_code_block = False
code_lines = []
line_count = 0

for line in lines:
    line_count += 1
    if line_count % 100 == 0:
        print(f"Processing line {line_count}/{len(lines)}...")
    
    # Handle code blocks
    if line.strip().startswith('```'):
        if in_code_block:
            # End of code block - add all code lines
            if code_lines:
                code_para = doc.add_paragraph()
                code_para.paragraph_format.left_indent = Inches(0.5)
                code_run = code_para.add_run('\n'.join(code_lines))
                code_run.font.name = 'Consolas'
                code_run.font.size = Pt(9)
                code_run.font.color.rgb = RGBColor(0, 0, 0)
                code_lines = []
            in_code_block = False
        else:
            in_code_block = True
        continue
    
    if in_code_block:
        code_lines.append(line.rstrip())
        continue
    
    # Skip empty lines
    if not line.strip():
        doc.add_paragraph()
        continue
    
    # Main title (# )
    if line.startswith('# ') and not line.startswith('## '):
        title = line[2:].strip()
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 153)  # Dark Blue
            run.font.size = Pt(28)
            run.bold = True
        continue
    
    # Subtitle (## with italic)
    if line.startswith('## ') and line.strip().startswith('## *'):
        subtitle = line[3:].strip().strip('*')
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(subtitle)
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(102, 102, 102)  # Gray
        run.italic = True
        continue
    
    # Section headings (## )
    if line.startswith('## ') and not line.startswith('### '):
        section = line[3:].strip()
        # Add some space before section
        doc.add_paragraph()
        heading = doc.add_heading(section, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.color.rgb = RGBColor(204, 0, 0)  # Red
            run.font.size = Pt(20)
            run.bold = True
        continue
    
    # Subsection headings (### )
    if line.startswith('### ') and not line.startswith('#### '):
        subsection = line[4:].strip()
        heading = doc.add_heading(subsection, level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 102, 204)  # Blue
            run.font.size = Pt(16)
            run.bold = True
        continue
    
    # Sub-subsection headings (#### )
    if line.startswith('#### '):
        subsubsection = line[5:].strip()
        heading = doc.add_heading(subsubsection, level=3)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 153, 76)  # Green
            run.font.size = Pt(14)
        continue
    
    # Handle tables
    if line.strip().startswith('|'):
        # Just add as formatted text for simplicity
        para = doc.add_paragraph(line.strip())
        para.paragraph_format.left_indent = Inches(0.3)
        for run in para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        continue
    
    # Handle horizontal rules
    if line.strip() == '---':
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run('━' * 60)
        run.font.color.rgb = RGBColor(0, 102, 204)
        continue
    
    # Regular paragraphs
    text = line.strip()
    para = doc.add_paragraph()
    
    # Handle bold text and emojis
    parts = text.split('**')
    for i, part in enumerate(parts):
        if i % 2 == 0:
            # Check for special formatting
            if part.startswith('✅') or part.startswith('📞') or part.startswith('📧') or \
               part.startswith('🌐') or part.startswith('📱') or part.startswith('📍') or \
               part.startswith('🕐') or part.startswith('💰') or part.startswith('🎁') or \
               part.startswith('📅') or part.startswith('✍️') or part.startswith('🎓') or \
               part.startswith('🚀'):
                run = para.add_run(part)
                run.font.size = Pt(12)
            else:
                para.add_run(part)
        else:
            run = para.add_run(part)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 51, 153)

# Save the document
print("\nSaving document as Neuro_Edge_Technologies_Brochure.docx...")
doc.save('Neuro_Edge_Technologies_Brochure.docx')

print(f"✅ Successfully created Neuro_Edge_Technologies_Brochure.docx")
print(f"📄 Professional brochure ready for printing and distribution!")

