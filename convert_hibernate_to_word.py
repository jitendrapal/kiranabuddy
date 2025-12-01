from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Read the markdown file
print("Reading Hibernate_Complete_Course.md...")
with open('Hibernate_Complete_Course.md', 'r', encoding='utf-8') as f:
    lines = f.readlines()

print(f"Processing {len(lines)} lines...")

in_code_block = False
code_lines = []
line_count = 0

for line in lines:
    line_count += 1
    if line_count % 500 == 0:
        print(f"Processing line {line_count}/{len(lines)}...")
    
    # Handle code blocks
    if line.strip().startswith('```'):
        if in_code_block:
            # End of code block - add all code lines
            if code_lines:
                code_para = doc.add_paragraph()
                code_para.paragraph_format.left_indent = Inches(0.5)
                code_run = code_para.add_run('\n'.join(code_lines))
                code_run.font.name = 'Consolas'
                code_run.font.size = Pt(9)
                code_run.font.color.rgb = RGBColor(0, 0, 0)
                code_lines = []
            in_code_block = False
        else:
            in_code_block = True
        continue
    
    if in_code_block:
        code_lines.append(line.rstrip())
        continue
    
    # Skip empty lines
    if not line.strip():
        doc.add_paragraph()
        continue
    
    # Main title (# )
    if line.startswith('# ') and not line.startswith('## '):
        title = line[2:].strip()
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 51, 153)
            run.font.size = Pt(24)
        continue
    
    # Chapter headings (## )
    if line.startswith('## ') and not line.startswith('### '):
        chapter = line[3:].strip()
        # Add page break before new chapter (except first)
        if line_count > 10:
            doc.add_page_break()
        heading = doc.add_heading(chapter, level=1)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(204, 0, 0)
            run.font.size = Pt(18)
        continue
    
    # Section headings (### )
    if line.startswith('### ') and not line.startswith('#### '):
        section = line[4:].strip()
        heading = doc.add_heading(section, level=2)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 102, 204)
            run.font.size = Pt(14)
        continue
    
    # Subsection headings (#### )
    if line.startswith('#### '):
        subsection = line[5:].strip()
        heading = doc.add_heading(subsection, level=3)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(0, 153, 76)
            run.font.size = Pt(12)
        continue
    
    # Handle tables
    if line.strip().startswith('|'):
        # Just add as formatted text for simplicity
        para = doc.add_paragraph(line.strip())
        para.paragraph_format.left_indent = Inches(0.3)
        for run in para.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        continue
    
    # Handle horizontal rules
    if line.strip() == '---':
        doc.add_paragraph('_' * 80)
        continue
    
    # Regular paragraphs
    text = line.strip()
    para = doc.add_paragraph()
    
    # Handle bold text
    parts = text.split('**')
    for i, part in enumerate(parts):
        if i % 2 == 0:
            para.add_run(part)
        else:
            run = para.add_run(part)
            run.bold = True

# Save the document
print("\nSaving document as Hibernate_Complete_Course.docx...")
doc.save('Hibernate_Complete_Course.docx')

# Calculate approximate pages
total_pages = len(lines) // 40
print(f"✅ Successfully created Hibernate_Complete_Course.docx")
print(f"📄 Total pages: approximately {total_pages}")

