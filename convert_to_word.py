"""
Convert Core Java Complete Book from Markdown to Word Document
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_word_document():
    """Convert markdown to Word document with proper formatting"""
    
    # Create a new Document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Read the markdown file
    print("Reading CORE_JAVA_COMPLETE_BOOK.md...")
    with open('CORE_JAVA_COMPLETE_BOOK.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    
    print(f"Processing {len(lines)} lines...")
    
    in_code_block = False
    code_lines = []
    
    for i, line in enumerate(lines):
        if i % 1000 == 0:
            print(f"Processing line {i}/{len(lines)}...")
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block - add all code lines
                if code_lines:
                    code_text = '\n'.join(code_lines)
                    p = doc.add_paragraph(code_text)
                    p.style = 'Normal'
                    for run in p.runs:
                        run.font.name = 'Consolas'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0, 0, 128)
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                code_lines = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            continue
        
        if in_code_block:
            code_lines.append(line)
            continue
        
        # Skip empty lines
        if not line.strip():
            doc.add_paragraph()
            continue
        
        # Main title (# with emoji)
        if line.startswith('# 📚'):
            p = doc.add_heading(line.replace('# 📚 ', ''), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(24)
                run.font.color.rgb = RGBColor(0, 51, 102)
                run.bold = True
            continue
        
        # Subtitle
        if line.startswith('## From Basics'):
            p = doc.add_paragraph(line.replace('## ', ''))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(102, 102, 102)
                run.italic = True
            continue
        
        # Chapter headings (## Chapter)
        if line.startswith('## Chapter'):
            doc.add_page_break()
            p = doc.add_heading(line.replace('## ', ''), level=1)
            for run in p.runs:
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0, 102, 204)
            continue
        
        # Section headings (## with emoji)
        if line.startswith('## 📖') or line.startswith('## 🎓'):
            p = doc.add_heading(line.replace('## ', '').replace('📖 ', '').replace('🎓 ', ''), level=1)
            for run in p.runs:
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0, 102, 204)
            continue
        
        # Other level 2 headings
        if line.startswith('## '):
            p = doc.add_heading(line.replace('## ', ''), level=2)
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(51, 102, 153)
            continue
        
        # Section headings (### with emoji)
        if line.startswith('### 📘') or line.startswith('### 💻') or line.startswith('### 🎯') or line.startswith('### 📝'):
            emoji = line.split()[0].replace('#', '').strip()
            text = line.replace('### ', '').replace(emoji, '').strip()
            p = doc.add_heading(text, level=2)
            for run in p.runs:
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(0, 128, 0)
                run.bold = True
            continue
        
        # Other level 3 headings
        if line.startswith('### '):
            p = doc.add_heading(line.replace('### ', ''), level=3)
            for run in p.runs:
                run.font.size = Pt(12)
            continue
        
        # Level 4 headings (####)
        if line.startswith('#### '):
            p = doc.add_heading(line.replace('#### ', ''), level=4)
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(102, 51, 0)
            continue
        
        # Horizontal rules
        if line.strip() == '---':
            p = doc.add_paragraph('_' * 80)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.color.rgb = RGBColor(192, 192, 192)
            continue
        
        # Bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Check for checkmarks
            if text.startswith('✅'):
                text = text.replace('✅', '').strip()
            p = doc.add_paragraph(text, style='List Bullet')
            continue
        
        # Numbered lists
        if re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            p = doc.add_paragraph(text, style='List Number')
            continue
        
        # Tables (simple markdown tables)
        if '|' in line and not line.strip().startswith('//'):
            # Skip table separator lines
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                continue
            # This is a table row - we'll handle it simply as text for now
            p = doc.add_paragraph(line.strip())
            for run in p.runs:
                run.font.size = Pt(10)
            continue
        
        # Regular paragraphs
        # Handle bold text **text**
        p = doc.add_paragraph()
        
        # Simple bold/italic handling
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = p.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(128, 0, 0)
            else:
                p.add_run(part)
    
    # Save the document
    output_file = 'CORE_JAVA_COMPLETE_BOOK.docx'
    print(f"\nSaving document as {output_file}...")
    doc.save(output_file)
    print(f"✅ Successfully created {output_file}")
    print(f"📄 Total pages: approximately {len(lines) // 40}")

if __name__ == '__main__':
    try:
        create_word_document()
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()

